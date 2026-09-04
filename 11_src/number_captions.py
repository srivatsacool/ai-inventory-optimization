"""Prefix figure/table captions with LaTeX numbers from report.aux.
Matches by document order + fuzzy caption text. Formatting-only."""
import os
import re
from docx import Document

ROOT = r'D:\Brain\07_Research\AI-INVENTORY-OPTIMIZATION'
FIN = os.path.join(ROOT, '09_reports', 'final')
DST = os.path.join(FIN, 'build_docx')
DOCX = os.path.join(FIN, 'research_report.docx')

aux = open(os.path.join(FIN, 'report.aux'), encoding='utf-8',
           errors='replace').read()
nums = dict(re.findall(r'\\newlabel\{([^}]+)\}\{\{([^}]*)\}', aux))

main = open(os.path.join(DST, 'report.tex'), encoding='utf-8').read()
order = re.findall(r'\\(?:input|include)\{([^}]+)\}', main)
envs = {'figure': [], 'table': []}
for rel in order:
    p = os.path.join(DST, rel if rel.endswith('.tex') else rel + '.tex')
    if not os.path.exists(p):
        continue
    t = open(p, encoding='utf-8').read()
    for m in re.finditer(r'\\begin\{(figure|table)\*?\}(.*?)\\end\{\1\*?\}',
                         t, re.DOTALL):
        kind, body = m.group(1), m.group(2)
        cm = re.search(r'\\caption\{(.*)\}\s*\\label\{([^}]+)\}', body, re.DOTALL)
        if not cm:
            cm = re.search(r'\\label\{([^}]+)\}', body)
            if not cm:
                continue
            envs[kind].append((cm.group(1), ''))
            continue
        cap, lab = cm.group(1), cm.group(2)
        cap = re.sub(r'\s+', ' ', re.sub(r'\\[a-zA-Z]+\{?', '', cap)).strip()
        envs[kind].append((lab, cap[:60]))

print('tex figs:', len(envs['figure']), 'tabs:', len(envs['table']))

d = Document(DOCX)
fig_caps = [p for p in d.paragraphs if p.style.name == 'Image Caption']
print('docx Image Captions:', len(fig_caps))
nf = 0
for p, (lab, cap) in zip(fig_caps, envs['figure']):
    n = nums.get(lab, '??')
    if not re.match(r'^Figure \d+:', p.text.strip()):
        p.text = f'Figure {n}: ' + p.text
        nf += 1
    elif cap and cap[:25].lower() not in p.text.lower():
        print('FIG MISMATCH:', lab, '|', p.text[:60])

nt = 0
ti = 0
for t in d.tables:
    prev = t._tbl.getprevious()
    if prev is None or not prev.tag.endswith('}p'):
        continue
    from docx.text.paragraph import Paragraph as P
    pp = P(prev, d)
    if ti >= len(envs['table']):
        break
    lab, cap = envs['table'][ti]
    txt = pp.text.strip()
    if lab and cap and cap[:25].lower() in txt.lower() \
            and not re.match(r'^Table \d+:', txt):
        pp.text = f"Table {nums.get(lab, '??')}: " + pp.text
        nt += 1
        ti += 1
    elif txt and len(txt) > 20 and pp.style.name not in (
            'Heading 1', 'Heading 2', 'Heading 3', 'Title'):
        # possible caption without fuzzy match; only consume if long prose
        if cap and cap[:25].lower() in txt.lower():
            pp.text = f"Table {nums.get(lab, '??')}: " + pp.text
            nt += 1
        ti += 1
print(f'numbered figs={nf} tables={nt}')
d.save(DOCX)
