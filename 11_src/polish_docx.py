"""Polish pandoc-built research_report.docx: native academic Word styling.
Formatting-only. Never touches research content."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = (r'D:\Brain\07_Research\AI-INVENTORY-OPTIMIZATION'
       r'\09_reports\final\research_report.docx')

d = Document(SRC)

# --- page setup: A4, 2.5cm margins (match LaTeX master) ---
for s in d.sections:
    s.page_width, s.page_height = Cm(21.0), Cm(29.7)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.5)
    s.header_distance, s.footer_distance = Cm(1.25), Cm(1.25)

# --- base styles ---
normal = d.styles['Normal']
normal.font.name, normal.font.size = 'Calibri', Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for i, (nm, sz) in enumerate([('Heading 1', 16), ('Heading 2', 13),
                                ('Heading 3', 11)], start=1):
    st = d.styles[nm]
    st.font.name, st.font.size = 'Calibri', Pt(sz)
    st.font.bold, st.font.color.rgb = True, RGBColor(0x1F, 0x1F, 0x1F)
    st.paragraph_format.space_before, st.paragraph_format.space_after = Pt(sz), Pt(4)

try:
    cap = d.styles['Caption']
    cap.font.name, cap.font.size, cap.font.italic = 'Calibri', Pt(9), True
    cap.paragraph_format.space_after = Pt(8)
except KeyError:
    pass

# --- title block: first 3 short paragraphs -> Title/Subtitle ---
titled = 0
for p in d.paragraphs[:8]:
    if titled >= 3 or p.style.name.startswith('Heading'):
        break
    if p.text.strip():
        p.style = d.styles['Title'] if titled == 0 else d.styles['Subtitle']
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titled += 1

# --- captions: paragraphs starting 'Figure N:' / 'Table N:' ---
ncap = 0
for p in d.paragraphs:
    if re.match(r'^(Figure|Table) \d+:', p.text.strip()):
        try:
            p.style = d.styles['Caption']
        except KeyError:
            for r in p.runs:
                r.italic, r.font.size = True, Pt(9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ncap += 1

# --- figures: center image paragraphs, cap width 6in ---
nimg = 0
for p in d.paragraphs:
    if p._p.xpath('.//w:drawing') or p._p.xpath('.//w:pict'):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            for el in r._r.xpath('.//wp:extent'):
                cx, cy = int(el.get('cx')), int(el.get('cy'))
                max_cx = int(Inches(6.0).emu)
                if cx > max_cx:
                    el.set('cx', str(max_cx))
                    el.set('cy', str(int(cy * max_cx / cx)))
        nimg += 1

# --- tables: grid style, header shading, compact font ---
ntab = 0
for t in d.tables:
    ntab += 1
    try:
        t.style = 'Table Grid'
    except KeyError:
        pass
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(t.rows):
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if ri == 0:
                        r.bold = True
            if ri == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'EDEDED')
                shd.set(qn('w:val'), 'clear')
                c._tc.get_or_add_tcPr().append(shd)

# --- header: short title; footer: page numbers ---
nhf = 0
for s in d.sections:
    h = s.header.paragraphs[0] if s.header.paragraphs else s.header.add_paragraph()
    if 'Forecasts to Shelf' not in h.text:
        h.text = 'From Forecasts to Shelf Decisions'
        for r in h.runs:
            r.font.size, r.italic = Pt(9), True
    f = s.footer.paragraphs[0] if s.footer.paragraphs else s.footer.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if 'PAGE' not in f._p.xml:
        f.text = ''
        r = f.add_run()
        f1 = OxmlElement('w:fldChar')
        f1.set(qn('w:fldCharType'), 'begin')
        r._r.append(f1)
        r2 = f.add_run()
        i2 = OxmlElement('w:instrText')
        i2.set(qn('xml:space'), 'preserve')
        i2.text = 'PAGE'
        r2._r.append(i2)
        r3 = f.add_run()
        f3 = OxmlElement('w:fldChar')
        f3.set(qn('w:fldCharType'), 'end')
        r3._r.append(f3)
    nhf += 1

# --- TOC field before first Heading 1 ---
nhead1 = sum(1 for p in d.paragraphs if p.style.name == 'Heading 1')
first_h1 = next((p for p in d.paragraphs if p.style.name == 'Heading 1'), None)
if first_h1 is not None and 'Table of Contents' not in d.paragraphs[0].text:
    toc_h = OxmlElement('w:p')
    first_h1._p.addprevious(toc_h)
    from docx.text.paragraph import Paragraph as P
    hp = P(toc_h, d)
    hp.style = d.styles['Heading 1']
    hp.add_run('Contents').bold = True
    fld_p = OxmlElement('w:p')
    toc_h.addnext(fld_p)
    fp = P(fld_p, d)
    r = fp.add_run()
    b = OxmlElement('w:fldChar')
    b.set(qn('w:fldCharType'), 'begin')
    r._r.append(b)
    ri = fp.add_run()
    it = OxmlElement('w:instrText')
    it.set(qn('xml:space'), 'preserve')
    it.text = r'TOC \o "1-3" \h \z \u'
    ri._r.append(it)
    rs = fp.add_run()
    sp = OxmlElement('w:fldChar')
    sp.set(qn('w:fldCharType'), 'separate')
    rs._r.append(sp)
    re_ = fp.add_run('Right-click > Update Field in Word to populate.')
    re_.italic = True
    rf = fp.add_run()
    e = OxmlElement('w:fldChar')
    e.set(qn('w:fldCharType'), 'end')
    rf._r.append(e)

cp = d.core_properties
cp.title = ('From Forecasts to Shelf Decisions: Evaluating AI-Based '
            'Inventory Optimization from Traditional Methods to Neural Forecasting')
cp.author = 'AI Inventory Optimization Research Project'

d.save(SRC)
print(f'captions={ncap} img_paras={nimg} tables={ntab} h1={nhead1} sections={nhf}')
print('saved ok')
