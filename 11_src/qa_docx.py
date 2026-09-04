import docx

d = docx.Document('research_report_draft.docx')
full = '\n'.join(p.text for p in d.paragraphs)
for t in d.tables:
    for row in t.rows:
        for c in row.cells:
            full += '\n' + c.text
print('leak_cref:', full.count('cref{'))
print('leak_label:', full.count('label{'))
print('leak_texttt:', full.count('texttt'))
print('leak_backslash_cmd:', sum(full.count(s) for s in ['\\times', '\\approx', '\\rightarrow', '\\%']))
print('leak_dollar:', full.count('$'))
print('qmarks:', full.count('[?]'))
print('chars:', len(full))
for n in ['1.316', '152.83', '2,084.49', '22.65', '25/27', '2,382', '0.978', '2282', '2282.31']:
    print(n, full.count(n))
heads = [p.text[:70] for p in d.paragraphs if p.style.name.startswith('Heading 1')]
print('H1 heads:')
for h in heads:
    print(' -', h)
